import io
import logging
from pathlib import Path
from typing import Any, Dict, List
from urllib.parse import urlparse

logger = logging.getLogger(__name__)


# ── Text extraction from file paths ──────────────────────────────────────────

def _load_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="ignore") as f:
        return f.read()


def _load_pdf(path: str) -> str:
    try:
        import pypdf
        parts = []
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except ImportError:
        logger.warning("pypdf not installed — skipping PDF: %s", path)
        return ""


def _load_docx(path: str) -> str:
    try:
        import docx
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — skipping DOCX: %s", path)
        return ""


def _load_epub(path: str) -> str:
    with open(path, "rb") as f:
        return _load_epub_bytes(f.read())


# ── Text extraction from bytes (for Streamlit uploads) ───────────────────────

def _load_pdf_bytes(content: bytes) -> str:
    try:
        import pypdf
        parts = []
        reader = pypdf.PdfReader(io.BytesIO(content))
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except ImportError:
        logger.warning("pypdf not installed — cannot read PDF bytes")
        return ""


def _load_docx_bytes(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    except ImportError:
        logger.warning("python-docx not installed — cannot read DOCX bytes")
        return ""


def _load_epub_bytes(content: bytes) -> str:
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(io.BytesIO(content))
        parts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "lxml")
            parts.append(soup.get_text(separator="\n", strip=True))
        return "\n\n".join(parts)
    except ImportError:
        logger.warning("ebooklib / beautifulsoup4 not installed — cannot read EPUB")
        return ""


# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Split text into overlapping character-level windows."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start += chunk_size - overlap
    return chunks


# ── Public: bytes upload (Streamlit) ─────────────────────────────────────────

def load_from_uploaded_bytes(
    filename: str,
    content: bytes,
    chunk_size: int = 512,
    overlap: int = 64,
) -> List[Dict[str, Any]]:
    """Process a Streamlit UploadedFile (raw bytes) into chunk dicts."""
    ext = Path(filename).suffix.lower()
    if ext in (".txt", ".md"):
        text = content.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        text = _load_pdf_bytes(content)
    elif ext == ".docx":
        text = _load_docx_bytes(content)
    elif ext == ".epub":
        text = _load_epub_bytes(content)
    else:
        logger.warning("Unsupported file type: %s", filename)
        return []

    chunks = chunk_text(text, chunk_size, overlap)
    stem = Path(filename).stem
    return [
        {"id": f"{stem}_chunk_{i}", "text": c, "source": filename, "chunk_index": i}
        for i, c in enumerate(chunks)
    ]


# ── Public: URL loading ───────────────────────────────────────────────────────

def load_from_url(
    url: str,
    chunk_size: int = 512,
    overlap: int = 64,
) -> List[Dict[str, Any]]:
    """Fetch a web page and extract its readable text into chunk dicts."""
    try:
        import requests
        from bs4 import BeautifulSoup

        resp = requests.get(
            url,
            timeout=20,
            headers={"User-Agent": "Mozilla/5.0 (AutoRAG document loader)"},
        )
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
    except Exception as exc:
        logger.error("Failed to fetch URL %s: %s", url, exc)
        return []

    if not text.strip():
        logger.warning("No text extracted from URL: %s", url)
        return []

    chunks = chunk_text(text, chunk_size, overlap)
    parsed = urlparse(url)
    stem = (parsed.netloc + parsed.path).replace("/", "_").replace(".", "_")[:60] or "url"
    return [
        {"id": f"{stem}_chunk_{i}", "text": c, "source": url, "chunk_index": i}
        for i, c in enumerate(chunks)
    ]


# ── Public: directory / file path loading ────────────────────────────────────

def load_documents(
    path: str,
    chunk_size: int = 512,
    overlap: int = 64,
) -> List[Dict[str, Any]]:
    """
    Recursively load .txt, .md, .pdf, .docx, .epub files from a path (file or dir).
    Returns a list of chunk dicts: {id, text, source, chunk_index}.
    """
    p = Path(path)
    if p.is_file():
        files = [p]
    elif p.is_dir():
        files = (
            list(p.rglob("*.txt"))
            + list(p.rglob("*.md"))
            + list(p.rglob("*.pdf"))
            + list(p.rglob("*.docx"))
            + list(p.rglob("*.epub"))
        )
    else:
        logger.error("Path does not exist: %s", path)
        return []

    docs: List[Dict[str, Any]] = []
    for f in files:
        try:
            ext = f.suffix.lower()
            if ext in (".txt", ".md"):
                text = _load_text(str(f))
            elif ext == ".pdf":
                text = _load_pdf(str(f))
            elif ext == ".docx":
                text = _load_docx(str(f))
            elif ext == ".epub":
                text = _load_epub(str(f))
            else:
                continue
            chunks = chunk_text(text, chunk_size, overlap)
            for i, chunk in enumerate(chunks):
                docs.append({
                    "id": f"{f.stem}_chunk_{i}",
                    "text": chunk,
                    "source": str(f),
                    "chunk_index": i,
                })
            logger.info("Loaded %d chunks from %s", len(chunks), f.name)
        except Exception as exc:
            logger.error("Failed to load %s: %s", f, exc)
    return docs
